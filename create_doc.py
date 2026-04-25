from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_presentation_script():
    doc = Document()
    
    # Title
    title = doc.add_heading('Presentation Script: Dynamic RAG-Based Customer Support Assistant', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Estimated Duration: 10-12 Minutes').bold = True
    doc.add_paragraph('Format: Live Presentation + Demo').bold = True
    doc.add_paragraph()
    
    def add_section(heading, duration, slide_text, body_text):
        doc.add_heading(f"{heading} ({duration})", level=1)
        p = doc.add_paragraph()
        p.add_run(f"[{slide_text}]").bold = True
        p.add_run("\n\n")
        p.add_run(body_text)
        doc.add_paragraph()

    # Section 1
    add_section(
        "1. Introduction", "1 Minute",
        "Slide/Screen: Title Slide with Project Name and Your Name",
        "\"Hello everyone, my name is [Your Name] from [Your Batch/Class]. Today, I am excited to present my project: a Dynamic RAG-Based Customer Support Assistant.\n\n"
        "The Problem: Traditional customer support chatbots are often rigid. They rely on pre-programmed scripts or static datasets, meaning they cannot adapt instantly when a company updates its policies, handbooks, or billing terms. Furthermore, when standard AI models face questions outside their training data, they tend to 'hallucinate' or guess the answers, which is unacceptable in a professional support environment.\n\n"
        "The Solution: I built an enterprise-grade AI assistant that dynamically learns from documents uploaded at runtime. It uses a Retrieval-Augmented Generation (RAG) architecture powered by LangGraph to route queries, extract context, and guarantee that the AI only answers based on the provided documents. If a query is too complex or legally sensitive, the system automatically escalates it to a human agent.\""
    )
    
    # Section 2
    add_section(
        "2. System Overview (High-Level Design)", "2 Minutes",
        "Slide/Screen: Architecture Diagram",
        "\"Let's look at how the system is built. The architecture is modular and highly scalable.\n\n"
        "Instead of just dumping data into a basic script, I engineered a complete pipeline with several key components interacting smoothly:\n\n"
        "1. Document Ingestion & Processing: When a user uploads a PDF, the system extracts the raw text. We use PyPDF for robust text extraction.\n"
        "2. Chunking & Embeddings: Because AI models have context limits, we can't feed them a 100-page manual all at once. The text is passed to our Chunker, which breaks it down into manageable pieces. These pieces are then converted into numerical vectors using a local Hugging Face embedding model.\n"
        "3. Dynamic Vector Database (ChromaDB): These vectors are stored in ChromaDB. What makes this system unique is its isolation: it dynamically creates a bespoke, isolated collection for *each session*. This ensures that User A's questions never accidentally pull data from User B's private documents.\n"
        "4. LangGraph Workflow Orchestration: This is the brain of the system. I used LangGraph to move away from linear chains. LangGraph allows the system to make decisions—evaluating if a query is relevant, if it requires clarification, or if it needs escalation.\n"
        "5. The LLM Layer: Finally, the Groq API (using Llama 3) receives the exact relevant chunks and generates a highly accurate, cited response.\""
    )
    
    # Section 3
    add_section(
        "3. End-to-End Workflow", "2.5 Minutes",
        "Slide/Screen: Workflow Flowchart / LangGraph Nodes",
        "\"So, what exactly happens when a user types a message? Let's trace a query through the system.\n\n"
        "1. User Query & Intent Routing: The user asks a question. The system doesn't immediately search the database. First, the LangGraph route_query node analyzes the intent. Is it a greeting? Is it a request to speak to a human? Or is it a factual question about the documents?\n"
        "2. Context Retrieval: If it's a factual question, the query goes to the Retriever. The system converts the user's question into a vector and asks ChromaDB: 'Give me the top 8 chunks of text most mathematically similar to this question.'\n"
        "3. Context Evaluation: Here is a critical safety check. The evaluate_context node reviews the retrieved chunks. If the chunks do not contain the answer, the graph physically halts the process and routes to a 'Fallback' state, forcing the AI to say, 'I cannot find this in the documents,' effectively preventing hallucinations.\n"
        "4. Response Generation: If the context is sufficient, the LLM processes the chunks and generates a response, complete with citations of the sources it used.\n"
        "5. Human-in-the-Loop (HITL): If the initial router detects a high-risk intent—like a user threatening legal action or complaining about billing—the graph routes directly to the escalate node. This bypasses the AI completely, generates a pending ticket, and alerts a human reviewer via the dashboard.\""
    )
    
    # Section 4
    add_section(
        "4. Live Demonstration", "3 Minutes",
        "Screen: Switch to the live web application (http://localhost:5173)",
        "\"I will now demonstrate the system live. As you can see, we have a premium, custom-built React frontend running on a FastAPI backend.\n\n"
        "(Action: Drag and drop a sample PDF, e.g., 'Employee_Handbook.pdf', into the sidebar)\n\n"
        "First, I'm uploading a sample Employee Handbook. You can see the loading state as the FastAPI backend processes, chunks, and indexes the document into a dynamic ChromaDB session. Once indexed, it appears here in the workspace.\n\n"
        "(Action: Type a valid question, e.g., 'What is the PTO policy?')\n\n"
        "Let's ask a direct question: 'What is the PTO policy?'\n"
        "Notice the typing indicator. The system routed the query, retrieved the context, evaluated it, and generated this response. You can click 'View Sources' here to see exactly which chunks of the PDF it pulled the answer from.\n\n"
        "(Action: Type an unsupported question, e.g., 'Who won the Super Bowl last year?')\n\n"
        "Now, let's try to trick it. I'll ask: 'Who won the Super Bowl last year?'\n"
        "Because of the LangGraph evaluation node, the system refuses to hallucinate and correctly states that this information is not in the uploaded documents.\n\n"
        "(Action: Type an escalation question, e.g., 'I want to sue you for billing fraud.')\n\n"
        "Finally, let's simulate a frustrated customer: 'I want to sue you for billing fraud.'\n"
        "Notice that the AI did not try to answer this. It instantly detected the legal/billing intent, generated an Escalation Ticket, and you can see it pop up right here in the Human Reviewer dashboard on the sidebar. I can manually resolve it as a human agent.\""
    )
    
    # Section 5
    add_section(
        "5. Technical Decisions", "1.5 Minutes",
        "Slide/Screen: Technical Stack & Architecture Details",
        "\"I want to highlight a few critical technical decisions I made during architecture design:\n\n"
        "- Chunk Size (1000) & Overlap (200): I experimented with different sizes and found that 1000 characters with a 200-character overlap struck the perfect balance. It's large enough to capture complete thoughts and paragraphs, but small enough to ensure the LLM isn't overwhelmed with irrelevant noise. The 200-character overlap prevents sentences from being cut in half during ingestion.\n"
        "- Embedding Model (all-MiniLM-L6-v2): I chose this specific Hugging Face model because it runs locally, is extremely lightweight, and executes incredibly fast without sacrificing semantic accuracy. It avoids the latency and cost of cloud embedding APIs.\n"
        "- Vector DB (ChromaDB): ChromaDB was chosen because it allows for lightweight, on-the-fly local collections. This was crucial for my 'workspace isolation' feature, ensuring data privacy between sessions.\n"
        "- Use of LangGraph: Traditional LangChain Chains are too linear. By using LangGraph, I could build a state machine with conditional edges. This allowed me to build the 'Evaluation' and 'HITL Escalation' gates, making the system cognitive rather than just reactive.\""
    )
    
    # Section 6
    add_section(
        "6. Challenges & Learnings", "1 Minute",
        "Slide/Screen: Challenges and Learnings",
        "\"Building this wasn't without its hurdles.\n\n"
        "- Technical Challenge 1: Initially, the system suffered from hallucinations. If the vector search returned poorly matched chunks, the LLM would just invent an answer.\n"
        "  Solution: I solved this by engineering the evaluate_context node in LangGraph. I forced the LLM to do a boolean check (Yes/No) on whether the context actually contained the answer *before* generating the final output.\n"
        "- Technical Challenge 2: Migrating from a monolithic Streamlit app to a decoupled FastAPI and React architecture introduced CORS and state management complexities.\n"
        "  Solution: I implemented proper RESTful endpoints and managed React session states locally, tying the UI session ID to the dynamic ChromaDB collection.\n\n"
        "Key Learning: The biggest takeaway for me was that a good AI product is 20% the LLM and 80% the data pipeline and orchestration around it. Prompt engineering is not enough; you need deterministic routing like LangGraph to build trust in the system.\""
    )
    
    # Section 7
    add_section(
        "7. Conclusion", "30 Seconds",
        "Slide/Screen: Conclusion and Future Scope",
        "\"In summary, I have built a dynamic, hallucination-resistant RAG assistant that prioritizes data privacy and safe routing via a Human-in-the-Loop architecture. It bridges the gap between raw AI capabilities and strict enterprise requirements.\n\n"
        "For future improvements, I plan to integrate OCR to handle scanned images within PDFs, and migrate the local ChromaDB to a cloud instance like Pinecone for persistent, cross-device workspaces.\n\n"
        "Thank you for your time, and I am happy to take any questions.\""
    )
    
    doc.save('Presentation_Script_RAG_Updated.docx')

if __name__ == "__main__":
    create_presentation_script()
